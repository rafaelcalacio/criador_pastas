import os
from docx import Document

## Solicita ao usuário para inserir o nome da pasta
pasta = input(f'Digite o nome da pasta \n')

nome = input('Digite o nome dos cômodos separando por espaço: \n')
pasta2 = nome.split()
y=len(pasta2)+1
## pasta2 = ['SALA', 'VARANDA', 'COZINHA', 'BANHEIRO', 'QUARTO1', 'QUARTO2', 'SUITE', 'LAVABO', 'AREA_DE_SERVICO']

def cria_pasta(diretorio):
    if not os.path.exists(diretorio):
        os.makedirs(diretorio)
        print(f'Diretório: {diretorio} criado com sucesso!!!')

        # Cria as duas subpastas padrão
        subpasta1 = os.path.join(diretorio, 'FOTO')
        subpasta2 = os.path.join(diretorio, 'VIDEO')

        os.makedirs(subpasta1)
        os.makedirs(subpasta2)

        print(f'Subpastas {subpasta1} e {subpasta2} criadas com sucesso.')

         # Cria um arquivo .docx no diretório principal
        criar_arquivo_docx(diretorio) 

        for x in range(0,y):
            pasta_x = os.path.join(subpasta1, f'{pasta2[x]}')
            os.makedirs(pasta_x)
            print(f'{pasta_x[x]} criada com sucesso!!!')

    else:
        print(f'O diretório {diretorio} já existe.')

def criar_arquivo_docx(diretorio):
    doc = Document()
    doc.add_heading(f'{pasta}', level=1)
    doc.add_paragraph('Descritivo:')
    
    # Caminho completo para o arquivo .docx
    caminho_arquivo_docx = os.path.join(diretorio, f'DESCRITIVO {pasta}.docx')
    
    doc.save(caminho_arquivo_docx)
    
    print(f'Arquivo {caminho_arquivo_docx} criado com sucesso.')

# Especifica o diretório desejado
diretorio_destino = r'C:\Users\Elenita\OneDrive\Pictures' + f'/{pasta}'

# Chama a função para criar as pastas
cria_pasta(diretorio_destino)
